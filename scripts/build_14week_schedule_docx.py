from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("세부_일정_14주차_정리본.docx")

FONT = "맑은 고딕"
INK = "162033"
BLUE = "1E3A8A"
MID_BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "DCFCE7"
GREEN_TEXT = "166534"
ORANGE = "FFF7ED"
ORANGE_TEXT = "9A3412"
BORDER = "CBD5E1"

TABLE_WIDTH_DXA = 13824
TABLE_INDENT_DXA = 120
COLUMN_WIDTHS_DXA = [720, 1800, 6480, 3600, 1224]
COLUMN_WIDTHS_IN = [0.5, 1.25, 4.5, 2.5, 0.85]

SCHEDULE = [
    ("1주", "문제 탐색", "캡스톤 오리엔테이션을 진행하고 생성형 AI·소설 생성 분야의 문제와 구현 가능성을 조사하였다.", "아이디어 후보 목록", "완료"),
    ("2주", "팀 구성", "팀 구성과 역할 분담을 완료하고 협업 방식, 개발 환경, 세부 일정을 정리하였다.", "세부 일정.docx", "완료"),
    ("3주", "주제 확정", "로컬 LLM과 JEPA-inspired predictor를 결합한 한국어 장편소설 생성 시스템으로 주제를 확정하였다.", "프로젝트 주제·목표", "완료"),
    ("4주", "계획 수립", "요구사항, 전체 파이프라인, 중간 산출물, 위험요소와 검증 방법을 계획서에 정리하였다.", "계획서·구성도", "완료"),
    ("5주", "피드백 반영", "계획서 피드백을 반영하고 LLM 전체 fine-tuning 대신 작은 predictor만 학습하는 범위를 확정하였다.", "수정 계획서", "완료"),
    ("6주", "개발 시작", "Streamlit scaffold, YAML config, Ollama client, 데이터·체크포인트·보고서 저장 구조를 구현하였다.", "a220ce3, d1fac12", "완료"),
    ("7주", "환경/발표", "CUDA 학습 지원, Windows launcher, 학습 정보와 live pipeline dashboard를 구축하고 중간 발표를 진행하였다.", "2f8645a, acd49db", "완료"),
    ("8주", "LLM+RAG", "합성 서사 데이터 생성, 장문 메모리 세션, 캐시와 평가 파이프라인을 연결하였다.", "e7bc4cc, 0926264", "완료"),
    ("9주", "단편 테스트", "JSON 추출·schema 안정화, 장르 프리셋, beat card, 인물 이름 일관성 검사를 추가하여 단편 생성을 검증하였다.", "aa5801d, 8a9dbd5", "완료"),
    ("10주", "코드 고도화", "Ollama 장애 복구, 합성 데이터 다양성 계획, pipeline trace, 캐시 분리와 저장소 정리 기능을 보강하였다.", "5012dd3, 4b062d3", "완료"),
    ("11주", "JEPA 삽입", "Residual MLP predictor, retrieval baseline, planner diagnostics와 predictor metadata 저장 기능을 구현하였다.", "d0f335a, dbf8c65", "완료"),
    ("12주", "장편 생성", "섹션별 장문 생성, 5,000자·10,000자·사용자 지정 길이, 저장 파일 이어쓰기와 중간 저장을 구현하였다.", "0926264, b0ea768", "완료"),
    ("13주", "Hallucination", "Creative Hallucination + JEPA 모드와 창의 확장률, 목표 정렬도, hallucination risk 평가 지표를 추가하였다.", "b0ea768, cf5f712", "완료"),
    ("14주", "최종 고도화", "Story-memory RAG, KG/state ledger, portable bundle, 장편 일관성 및 섹션 대기시간을 개선하고 최종 보고서·시연 자료를 작성하였다.", "de64f3e~141fca3", "진행"),
]


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in COLUMN_WIDTHS_DXA:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width_dxa, width_in in zip(row.cells, COLUMN_WIDTHS_DXA, COLUMN_WIDTHS_IN):
            cell.width = Inches(width_in)
            _set_cell_width(cell, width_dxa)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 9.5, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_schedule_table(document: Document, rows: list[tuple[str, str, str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["주차", "단계", "주요 활동 및 구현", "주요 산출물 / Git 근거", "상태"]
    for idx, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
        _set_cell_shading(cell, BLUE)
        _set_cell_text(cell, text, bold=True, color="FFFFFF", size=10, center=idx in (0, 1, 4))
    _set_repeat_table_header(table.rows[0])

    for row_index, item in enumerate(rows):
        row = table.add_row()
        if row_index % 2:
            for cell in row.cells:
                _set_cell_shading(cell, LIGHT_GRAY)
        for col_index, (cell, text) in enumerate(zip(row.cells, item)):
            center = col_index in (0, 1, 4)
            bold = col_index in (0, 1, 4)
            color = INK
            if col_index == 4:
                if text == "완료":
                    _set_cell_shading(cell, GREEN)
                    color = GREEN_TEXT
                else:
                    _set_cell_shading(cell, ORANGE)
                    color = ORANGE_TEXT
            _set_cell_text(cell, text, bold=bold, color=color, center=center)

    _set_table_geometry(table)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, MID_BLUE, 18, 10),
        ("Heading 2", 13, MID_BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run("Novel JEPA Lab 캡스톤디자인 14주 개발 일정")
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("세부 일정 문서와 GitHub 커밋 이력 38건을 대조하여 재구성")
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("52606D")

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run("작성 기준: 2026년 6월 13일  |  개발 이력 범위: 2026년 5월 14일~6월 6일")
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("64748B")


def _add_basis_notes(document: Document) -> None:
    document.add_heading("일정 구성 근거", level=2)
    notes = [
        "원본 세부 일정.docx의 2~14주차 계획을 유지하고, 누락된 1주차를 오리엔테이션 및 문제 탐색 단계로 보완하였다.",
        "로컬 Git 저장소의 38개 커밋 메시지와 날짜를 검토하여 개발 시작 이후의 구현 순서를 실제 기능 단위로 재배치하였다.",
        "표의 Git 근거는 각 주차를 대표하는 커밋 해시이며, 해당 주차의 모든 세부 커밋을 나열한 것은 아니다.",
        "1~13주차는 완료, 14주차는 최종 보고서와 시연 자료를 보완하는 진행 단계로 정리하였다.",
    ]
    for note in notes:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(note)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(10.5)


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    _configure_styles(document)
    _add_title(document)

    document.add_heading("1~7주차: 기획·기반·개발환경", level=2)
    _add_schedule_table(document, SCHEDULE[:7])

    document.add_page_break()
    document.add_heading("8~14주차: 핵심 기능·고도화·최종화", level=2)
    _add_schedule_table(document, SCHEDULE[7:])
    _add_basis_notes(document)

    document.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
